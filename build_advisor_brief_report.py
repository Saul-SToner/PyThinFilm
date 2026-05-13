from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"C:\Users\L2791\Downloads\Vscode")
IMG_ROOT = ROOT / "advisor_image_bundle精选版_20260513"
OUT_DIR = ROOT / "advisor_reports"
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = OUT_DIR / "导师简略汇报_薄膜与表面结构专题.docx"


SECTIONS = [
    (
        "教学主树",
        "已完成单层减反膜、F-P 滤光片和高反膜的理论与 COMSOL 对照，形成了主树标准验证链。",
        [
            "01_教学主树_标准验证总览.png",
            "03_教学主树_F-P滤光片_主图.png",
            "04_教学主树_高反膜_主图.png",
        ],
    ),
    (
        "高级减反",
        "高级减反专题用于展示从单层减反到多孔层、再到等效渐变结构的演化关系。",
        [
            "01_高级减反_专题总览.png",
            "02_高级减反_结构家族图谱.png",
        ],
    ),
    (
        "多孔双层减反",
        "多孔二氧化硅双层减反结构已经完成验证、参数敏感性和角度稳定性分析，是当前减反家族中的强表现方案。",
        [
            "01_多孔双层减反_验证主图.png",
            "02_多孔双层减反_参数敏感性.png",
            "03_多孔双层减反_角度稳定性.png",
        ],
    ),
    (
        "蛾眼结构",
        "蛾眼结构专题已完成基准模型、几何参数扫描和最终推荐参数的整理，可用于展示亚波长表面结构减反路线。",
        [
            "01_蛾眼结构_最终反射谱对比.png",
            "02_蛾眼结构_高度扫描.png",
            "03_蛾眼结构_周期扫描.png",
        ],
    ),
    (
        "吸收表面",
        "吸收表面专题已完成平面基准、粗糙版本对比和吸收增益趋势分析，说明粗糙度增大会提升吸收增强效果。",
        [
            "03_吸收表面_平面与粗糙增益对比.png",
            "04_吸收表面_粗糙度趋势.png",
            "05_吸收表面_吸收增益趋势.png",
        ],
    ),
]


def set_run_font(run, east_asia: str = "Microsoft YaHei", latin: str = "Calibri"):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def style_normal(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)


def add_field(paragraph, instruction: str):
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), instruction)
    run = paragraph.add_run()
    run._r.append(fld_simple)
    return run


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "目录将在 Word 中更新。"
    fld_sep.append(fld_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, fld_end):
        run._r.append(el)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    run = p.add_run("薄膜与表面结构阶段性简略汇报")
    set_run_font(run)
    run.font.size = Pt(22)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("教学主树、高级减反、蛾眼结构与吸收表面专题")
    set_run_font(run2)
    run2.font.size = Pt(13)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(120)
    run3 = p3.add_run("自动整理版本")
    set_run_font(run3)
    run3.font.size = Pt(11)

    doc.add_page_break()


def add_summary(doc: Document):
    h = doc.add_paragraph(style="Heading 1")
    r = h.add_run("一、汇报目录")
    set_run_font(r)

    p = doc.add_paragraph()
    add_toc(p)
    doc.add_page_break()

    h2 = doc.add_paragraph(style="Heading 1")
    r2 = h2.add_run("二、总体结论")
    set_run_font(r2)
    bullets = [
        "教学主树的标准验证链已经打通，单层减反膜、F-P 滤光片和高反膜的理论与 COMSOL 曲线吻合良好。",
        "高级减反方向已形成从单层、多孔单层、多孔双层到蛾眼等效渐变层的结构家族。",
        "多孔双层减反结构当前表现最强，且已完成参数敏感性与角度稳定性分析。",
        "蛾眼结构已完成几何参数扫描并得到当前推荐参数，可作为亚波长表面减反代表对象。",
        "粗糙吸收表面相对平面基准表现出明确吸收增益，说明表面结构化对吸收增强有效。",
    ]
    for text in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        set_run_font(run)
    doc.add_page_break()


def add_topic_section(doc: Document, topic: str, summary: str, image_names: list[str]):
    heading = doc.add_paragraph(style="Heading 1")
    run = heading.add_run(topic)
    set_run_font(run)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(summary)
    set_run_font(run)

    base = IMG_ROOT / topic
    for idx, image_name in enumerate(image_names, start=1):
        image_path = base / image_name
        if not image_path.exists():
            continue
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f"{topic} 图 {idx}：{image_name.rsplit('.', 1)[0].split('_', 2)[-1]}")
        set_run_font(cap_run)
        cap_run.bold = True
        doc.add_picture(str(image_path), width=Cm(15.8))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)

    doc.add_page_break()


def build():
    doc = Document()
    style_normal(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_cover(doc)
    add_summary(doc)
    for topic, summary, image_names in SECTIONS:
        add_topic_section(doc, topic, summary, image_names)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
